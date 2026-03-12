import io
import re
import json
from pathlib import Path
import pdfplumber
from docx import Document
from docx.shared import Pt
from services.llm_client import get_client, TEXT_MODEL


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """Extract text from PDF or DOCX resume."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif ext in (".docx", ".doc"):
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(para.text for para in doc.paragraphs)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


async def optimize_resume(
    resume_text: str,
    job_data: dict,
) -> tuple[bytes, dict]:
    """
    Use Claude to optimize resume for the job.
    Returns (docx_bytes, metadata_dict).
    """

    job_title = job_data.get("title", "the role")
    company = job_data.get("company", "the company")
    job_desc = job_data.get("description", "")
    requirements = job_data.get("requirements", [])

    prompt = f"""You are an expert resume writer and ATS optimization specialist.

I need you to rewrite and optimize this resume for the following job.

**Target Job:** {job_title} at {company}

**Key Requirements:**
{chr(10).join(f'- {r}' for r in requirements[:10])}

**Full Job Description:**
{job_desc[:3000]}

**Original Resume:**
{resume_text}

**Instructions:**
1. Tailor the resume to match the job requirements naturally
2. Add relevant keywords from the job description throughout
3. Quantify achievements where possible (add estimates if none exist)
4. Reorder bullet points to put most relevant first
5. Strengthen action verbs
6. Ensure ATS compatibility (avoid tables, graphics descriptions)
7. Keep it truthful - enhance presentation but don't fabricate
8. Maintain the original structure/sections

Return a JSON object with:
{{
  "optimized_resume": "full optimized resume text with \\n for line breaks",
  "key_changes": ["change 1", "change 2", "change 3"],
  "keywords_added": ["keyword1", "keyword2", "keyword3"],
  "match_score": 8.5
}}"""

    response = get_client().chat.completions.create(
        model=TEXT_MODEL,
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}]
    )

    text = response.choices[0].message.content

    # Parse JSON response
    json_match = re.search(r'\{.*\}', text, re.DOTALL)
    if not json_match:
        raise ValueError("Could not parse optimizer response")

    data = json.loads(json_match.group())
    optimized_text = data["optimized_resume"]

    # Create DOCX
    docx_bytes = _create_resume_docx(optimized_text)

    metadata = {
        "job_title": job_title,
        "company": company,
        "match_score": data.get("match_score", 7.0),
        "key_changes": data.get("key_changes", []),
        "keywords_added": data.get("keywords_added", []),
    }

    return docx_bytes, metadata


def _create_resume_docx(text: str) -> bytes:
    """Convert optimized text back to a clean DOCX."""
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    lines = text.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Detect headers (ALL CAPS or lines followed by dashes)
        if line.isupper() and len(line) > 3:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)
        elif line.startswith("•") or line.startswith("-") or line.startswith("*"):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.lstrip("•-* "))
            p.paragraph_format.space_after = Pt(2)
        elif line.endswith(":") or (len(line) < 50 and line == line.title()):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            p.paragraph_format.space_before = Pt(6)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
