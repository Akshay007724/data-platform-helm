import io

from docx import Document

from .base import BaseExtractor, ExtractedContent, PageContent, TableContent


class WordExtractor(BaseExtractor):
    def extract(self, file_bytes: bytes) -> ExtractedContent:
        doc = Document(io.BytesIO(file_bytes))

        # Paragraphs (excludes table cells — python-docx separates them)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        # Tables
        tables: list[TableContent] = []
        for table in doc.tables:
            if not table.rows:
                continue
            rows_data = [
                [cell.text.strip() for cell in row.cells] for row in table.rows
            ]
            # Deduplicate merged cells (docx repeats merged cell text)
            deduped = []
            for row in rows_data:
                seen_in_row: set[str] = set()
                clean_row = []
                for cell in row:
                    if cell not in seen_in_row or not cell:
                        clean_row.append(cell)
                        seen_in_row.add(cell)
                    else:
                        clean_row.append("")
                deduped.append(clean_row)

            headers = deduped[0] if deduped else []
            data_rows = deduped[1:] if len(deduped) > 1 else []
            if any(any(c for c in row) for row in [headers] + data_rows):
                tables.append(TableContent(page=1, headers=headers, rows=data_rows))

        # Images
        image_blobs: list[bytes] = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_blobs.append(rel.target_part.blob)

        page = PageContent(page=1, text=text, images=image_blobs, tables=tables)
        return ExtractedContent(pages=[page], file_type="word")
